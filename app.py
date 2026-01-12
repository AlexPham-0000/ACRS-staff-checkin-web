import os
from datetime import datetime

from flask import Flask, render_template, request, redirect, url_for, flash
from docx import Document
import os
from flask_sqlalchemy import SQLAlchemy
from sqlalchemy import func
import pandas as pd
from openpyxl import load_workbook
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
from zoneinfo import ZoneInfo
from datetime import datetime
import io
from flask import send_file
from datetime import time as dtime



SEA_TZ = ZoneInfo("America/Los_Angeles")

def now_seattle_naive():
    return datetime.now(SEA_TZ).replace(tzinfo=None)

def today_seattle():
    return now_seattle_naive().date()


app = Flask(__name__)

app.secret_key = "change-this-secret"
# --------- Cấu hình database (SQLite) ---------
db_url = os.environ.get("DATABASE_URL")

# Render may provide postgres:// but SQLAlchemy needs postgresql://
if db_url and db_url.startswith("postgres://"):
    db_url = db_url.replace("postgres://", "postgresql://", 1)

if db_url:
    app.config["SQLALCHEMY_DATABASE_URI"] = db_url
else:
    basedir = os.path.abspath(os.path.dirname(__file__))
    db_path = os.path.join(basedir, "staff_checkin.db")
    app.config["SQLALCHEMY_DATABASE_URI"] = "sqlite:///" + db_path

app.config["SQLALCHEMY_TRACK_MODIFICATIONS"] = False
db = SQLAlchemy(app)

# Đường dẫn đến file danh sách nhân viên
STAFF_DOCX_PATH = os.path.join(os.path.dirname(__file__), "ACRSstaff_name.docx")

STAFF_DOCX_PATH = "ACRSstaff_name.docx"



def add_staff_to_docx(new_name: str) -> bool:
    """Add a staff name to docx if not exists. Return True if added, False if duplicate."""
    new_name = (new_name or "").strip()
    if not new_name:
        return False

    # đọc danh sách hiện tại
    existing = load_staff_names_from_docx(STAFF_DOCX_PATH)
    if any(n.lower() == new_name.lower() for n in existing):
        return False

    # mở file nếu có, không thì tạo mới
    if os.path.exists(STAFF_DOCX_PATH):
        doc = Document(STAFF_DOCX_PATH)
    else:
        doc = Document()

    doc.add_paragraph(new_name)
    doc.save(STAFF_DOCX_PATH)

    return True



def load_staff_names_from_docx(path: str) -> list[str]:
    names = []

    # If missing OR empty/broken docx -> create a fresh docx
    if (not os.path.exists(path)) or (os.path.getsize(path) == 0):
        doc = Document()
        doc.save(path)
        return names

    try:
        doc = Document(path)
    except Exception:
        # If file exists but is not a valid .docx (corrupted / not synced)
        doc = Document()
        doc.save(path)
        return names

    for p in doc.paragraphs:
        if p.text and p.text.strip():
            names.append(p.text.strip())

    return names

# Danh sách tên được phép dùng
ALLOWED_STAFF_NAMES = load_staff_names_from_docx(STAFF_DOCX_PATH)
print("Loaded", len(ALLOWED_STAFF_NAMES), "staff names")   # để xem trên terminal

# --------- Danh sách phòng ban cố định ---------
DEPARTMENTS = [
    "AAS",
    "BH",
    "E&C",
    "CYF",
    "CE",
    "RS",
    "ADMIN",
    "ADMIN FACILITIES",
    "KITCHEN",
    "CLUB BAMBOO",
    "DATA",
    "ADMIN GENERAL",
    "HR",
    "IT",
    "FINANCE",
    "DEV",
    "GENOA",
]


# --------- Models (bảng trong DB) ---------
class Staff(db.Model):
    __tablename__ = "staff"

    id = db.Column(db.Integer, primary_key=True)
    name = db.Column(db.String(120), nullable=False)
    department = db.Column(db.String(120), nullable=False)


class CheckIn(db.Model):
    id = db.Column(db.Integer, primary_key=True)
    staff_id = db.Column(db.Integer, db.ForeignKey("staff.id"), nullable=False)
    time_in = db.Column(db.DateTime, nullable=False)
    time_out = db.Column(db.DateTime, nullable=True)
    note = db.Column(db.String(255))   # 🆕 Ghi chú
    returned_item = db.Column(db.Boolean, default=False, nullable=False)

    staff = db.relationship("Staff", backref=db.backref("checkins", lazy=True))


# Tạo bảng nếu chưa có
with app.app_context():
    db.create_all()



# --------- Trang chính: Today Check-ins ---------
@app.route("/", methods=["GET", "POST"])
def index():

    if request.method == "POST":
        name = (request.form.get("name") or request.form.get("staff_name") or "").strip()
        department = request.form.get("department", "").strip()
        note = request.form.get("note", "").strip()
        action = request.form.get("action")
        new_name = request.form.get("new_staff_name", "").strip()

        using_new = bool(new_name)
        using_existing = bool(name)

        # Must not fill both
        if using_new and using_existing:
            flash("Please use ONLY ONE: choose an existing staff OR type a new staff name (not both).")
            return redirect(url_for("index"))

        # Decide which name to use
        chosen_name = new_name if using_new else name

        # Validate action + department + name
        if action not in ("in", "out"):
            flash("Please select an action.")
            return redirect(url_for("index"))

        if not department:
            flash("Please select a department.")
            return redirect(url_for("index"))

        if not chosen_name:
            flash("Please choose an existing staff name or type a new one.")
            return redirect(url_for("index"))

        # ✅ Existing name must be in DOCX list
        if not using_new:
            allowed_lower = {n.lower() for n in ALLOWED_STAFF_NAMES}
            if chosen_name.lower() not in allowed_lower:
                flash("Name not found in ACRSstaff_name.docx. Use 'Add New Staff' box for new names.")
                return redirect(url_for("index"))

        canonical_name = chosen_name

        today = today_seattle()
        start = datetime.combine(today, datetime.min.time())
        end = datetime.combine(today, datetime.max.time())

        # get/create Staff for this name+department
        staff = Staff.query.filter_by(name=canonical_name, department=department).first()
        if staff is None:
            staff = Staff(name=canonical_name, department=department)
            db.session.add(staff)
            db.session.commit()

        # ---------- CHECK IN ----------
        if action == "in":
            # block check-in if already checked in today under another dept
            existing_today_any_dept = (
                CheckIn.query.join(Staff)
                .filter(
                    func.lower(Staff.name) == canonical_name.lower(),
                    CheckIn.time_in >= start,
                    CheckIn.time_in <= end,
                )
                .order_by(CheckIn.time_in.desc())
                .first()
            )
            if existing_today_any_dept:
                existing_dept = existing_today_any_dept.staff.department
                if existing_dept != department:
                    flash(f"{canonical_name} already checked in under department '{existing_dept}'.")
                    return redirect(url_for("index"))

            latest_today = (
                CheckIn.query
                .filter(
                    CheckIn.staff_id == staff.id,
                    CheckIn.time_in >= start,
                    CheckIn.time_in <= end,
                )
                .order_by(CheckIn.time_in.desc())
                .first()
            )
            if latest_today:
                if latest_today.time_out is None:
                    flash(f"{canonical_name} already checked in today and has not checked out yet.")
                else:
                    flash(f"{canonical_name} has already checked in and checked out today. No more check-ins allowed today.")
                return redirect(url_for("index"))

            ci = CheckIn(
                staff_id=staff.id,
                time_in=now_seattle_naive(),
                note=note,
                returned_item=False,
            )
            db.session.add(ci)
            db.session.commit()
            flash(f"{canonical_name} checked in.")
            return redirect(url_for("index"))

        # ---------- CHECK OUT ----------
        open_checkin = (
            CheckIn.query.join(Staff)
            .filter(
                func.lower(Staff.name) == canonical_name.lower(),
                Staff.department == department,
                CheckIn.time_out.is_(None),
                CheckIn.time_in >= start,
                CheckIn.time_in <= end,
            )
            .order_by(CheckIn.time_in.desc())
            .first()
        )

        if not open_checkin:
            any_open_other_dept = (
                CheckIn.query.join(Staff)
                .filter(
                    func.lower(Staff.name) == canonical_name.lower(),
                    CheckIn.time_out.is_(None),
                    CheckIn.time_in >= start,
                    CheckIn.time_in <= end,
                )
                .order_by(CheckIn.time_in.desc())
                .first()
            )
            if any_open_other_dept:
                existing_dept = any_open_other_dept.staff.department
                flash(f"{canonical_name} is checked in under department '{existing_dept}'. Please select that department to check out.")
                return redirect(url_for("index"))

            flash(f"{canonical_name} has not checked in today, so there is nothing to check out.")
            return redirect(url_for("index"))

        # block checkout if borrowed item noted but not returned
        if (open_checkin.note and open_checkin.note.strip()) and (not open_checkin.returned_item):
            flash(f"{canonical_name} hasn't returned: {open_checkin.note}")
            return redirect(url_for("index"))

        open_checkin.time_out = now_seattle_naive()
        db.session.commit()
        flash(f"{canonical_name} checked out.")
        return redirect(url_for("index"))

    # ---------- GET ----------
    today = today_seattle()
    start = datetime.combine(today, datetime.min.time())
    end = datetime.combine(today, datetime.max.time())

    records = (
        CheckIn.query.join(Staff)
        .filter(CheckIn.time_in >= start, CheckIn.time_in <= end)
        .order_by(
            Staff.department.asc(),
            Staff.name.asc(),
            CheckIn.time_in.asc(),
        )
        .all()
    )
    ## ✅ PIN NOTE STAFF TO TOP
    records = sorted(records, key=lambda r: 0 if (r.note and r.note.strip()) else 1)

# --------- Trang admin ---------
@app.route("/admin")
def admin():
    # danh sách staff
    staff_list = Staff.query.order_by(
        Staff.department.asc(),
        Staff.name.asc()
    ).all()

    print("DEBUG server datetime.today().date():", datetime.today().date())
    print("DEBUG seattle today_seattle():", today_seattle())
    # 100 record gần nhất, cũng nhóm theo department + name
    records = (
        CheckIn.query.join(Staff)
        .order_by(
            Staff.department.asc(),
            Staff.name.asc(),
            CheckIn.time_in.desc()
        )
        .limit(100)
        .all()
    )

    return render_template(
        "admin.html",
        staff_list=staff_list,
        records=records,
    )


import pandas as pd

@app.route("/delete/<int:record_id>", methods=["POST"])
def delete_record(record_id):
    record = CheckIn.query.get(record_id)

    if record:
        db.session.delete(record)
        db.session.commit()
        flash("Record deleted successfully.")
    else:
        flash("Record not found.")

    return redirect(url_for("index"))

@app.route("/export")
def export_excel():
    # ----- Lấy dữ liệu CHỈ CHO HÔM NAY -----
    today = today_seattle()
    start = datetime.combine(today, datetime.min.time())
    end = datetime.combine(today, datetime.max.time())

    today_records = (
        CheckIn.query.join(Staff)
        .filter(CheckIn.time_in >= start, CheckIn.time_in <= end)
        .order_by(
            Staff.department.asc(),
            Staff.name.asc(),
            CheckIn.time_in.asc()
        )
        .all()
    )

    if not today_records:
        return f"No records to export for {today}."

    rows = []

    # Duyệt từng department theo thứ tự trong DEPARTMENTS
    for dept in DEPARTMENTS:
        dept_records = [r for r in today_records if r.staff.department == dept]
        if not dept_records:
            continue

        # Dòng tiêu đề của department (sau này sẽ bold + nền xám)
        rows.append({
            "Name": "",
            "Note": "",
            "Returned": "",
            "Department": dept,
            "Date": "",
            "Time In": "",
            "Time Out": "",
        })

        # Các dòng nhân viên của phòng ban đó
        for r in dept_records:
            rows.append({
                "Name": r.staff.name,
                "Note": r.note or "",
                "Returned": ("Yes" if r.returned_item else "No") if (r.note and r.note.strip()) else "",
                "Department": r.staff.department,
                "Date": r.time_in.date(),
                "Time In": r.time_in.strftime("%I:%M %p"),
                "Time Out": r.time_out.strftime("%I:%M %p") if r.time_out else "",
            })

        # Dòng trống ngăn cách giữa các department
        rows.append({
            "Name": "",
            "Note": "",
            "Returned": "",
            "Department": "",
            "Date": "",
            "Time In": "",
            "Time Out": "",
        })

    # Tạo DataFrame với thứ tự cột rõ ràng
    df = pd.DataFrame(
        rows,
        columns=["Name", "Note", "Returned", "Department", "Date", "Time In", "Time Out"]
    )

    # ----- Tạo tên file: ACRScheckinMM-DD-YYYY_HHMMSS.xlsx -----
    date_str = today.strftime("%m-%d-%Y")
    time_str = now_seattle_naive().strftime("%H%M%S")
    filename = f"ACRScheckin{date_str}_{time_str}.xlsx"

    # ====== TẠO EXCEL TRONG BỘ NHỚ (KHÔNG LƯU FILE TRÊN RENDER) ======
    output = io.BytesIO()

    # Ghi df vào excel (memory)
    df.to_excel(output, index=False, engine="openpyxl")
    output.seek(0)

    # Load workbook từ memory để format như bạn đang làm
    wb = load_workbook(output)
    ws = wb.active

    center_align = Alignment(horizontal="center", vertical="center")
    bold_font = Font(bold=True)

    header_fill = PatternFill(fill_type="solid", fgColor="D9D9D9")    # xám cho header & dept
    note_fill = PatternFill(fill_type="solid", fgColor="FFF9C4")      # vàng nhạt cho Note

    thin_border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin"),
    )

    max_row = ws.max_row
    max_col = ws.max_column

    for row in ws.iter_rows(min_row=1, max_row=max_row, min_col=1, max_col=max_col):
        row_idx = row[0].row

        # canh giữa + border cho tất cả
        for cell in row:
            cell.alignment = center_align
            cell.border = thin_border

        # Hàng tiêu đề đầu tiên
        if row_idx == 1:
            for cell in row:
                cell.font = bold_font
                cell.fill = header_fill
            continue

        # detect department header rows (Name, Note, Returned, Date, Time In, Time Out trống)
        name_val = row[0].value
        note_val = row[1].value
        returned_val = row[2].value
        dept_val = row[3].value
        date_val = row[4].value
        tin_val = row[5].value
        tout_val = row[6].value

        if (
            (name_val in [None, ""]) and
            (note_val in [None, ""]) and
            (returned_val in [None, ""]) and
            (date_val in [None, ""]) and
            (tin_val in [None, ""]) and
            (tout_val in [None, ""]) and
            isinstance(dept_val, str) and dept_val in DEPARTMENTS
        ):
            for cell in row:
                cell.font = bold_font
                cell.fill = header_fill

    # 🟡 Tô vàng cột Note (col 2) nếu có nội dung
    for row_idx in range(2, max_row + 1):
        cell = ws.cell(row=row_idx, column=2)  # column 2 = Note
        if cell.value not in (None, ""):
            cell.fill = note_fill
            cell.alignment = Alignment(horizontal="left", vertical="center")

    # Save workbook formatted -> new memory buffer
    final_output = io.BytesIO()
    wb.save(final_output)
    final_output.seek(0)

    # ✅ TRẢ FILE VỀ BROWSER ĐỂ TẢI XUỐNG
    return send_file(
        final_output,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )


@app.route("/toggle_returned/<int:checkin_id>", methods=["POST"])
def toggle_returned(checkin_id):
    ci = CheckIn.query.get_or_404(checkin_id)
    ci.returned_item = not ci.returned_item
    db.session.commit()
    return redirect(url_for("index"))

@app.route("/add_staff_db", methods=["POST"])
def add_staff_db():
    name = request.form.get("new_staff_name", "").strip()
    department = request.form.get("department", "").strip()

    if not name or not department:
        flash("Please enter name and department.")
        return redirect(url_for("index"))

    today = today_seattle()
    start = datetime.combine(today, datetime.min.time())
    end = datetime.combine(today, datetime.max.time())

    # 1) get/create Staff
    staff = Staff.query.filter(
        func.lower(Staff.name) == name.lower(),
        Staff.department == department
    ).first()

    if staff is None:
        staff = Staff(name=name, department=department)
        db.session.add(staff)
        db.session.commit()

    # 2) block if already checked in today (same staff+dept)
    already_today = (
        CheckIn.query
        .filter(
            CheckIn.staff_id == staff.id,
            CheckIn.time_in >= start,
            CheckIn.time_in <= end,
        )
        .order_by(CheckIn.time_in.desc())
        .first()
    )
    if already_today:
        flash("This staff is already checked in today.")
        return redirect(url_for("index"))

    # 3) auto check-in now
    ci = CheckIn(
        staff_id=staff.id,
        time_in=now_seattle_naive(),
        note="",
        returned_item=False,
    )
    db.session.add(ci)
    db.session.commit()

    flash(f"Added & checked in: {name} ({department})")
    return redirect(url_for("index"))

from flask import send_from_directory

from flask import send_from_directory

@app.route("/service-worker.js")
def service_worker():
    return send_from_directory("static", "service-worker.js")

@app.route("/edit_time/<int:checkin_id>", methods=["POST"])
def edit_time(checkin_id):
    ci = CheckIn.query.get_or_404(checkin_id)

    which = (request.form.get("which") or "").strip().lower()   # "in" or "out"
    tstr = (request.form.get("time") or "").strip()             # "09:15" or "9:15 AM"

    if which not in ("in", "out"):
        flash("Invalid edit request.")
        return redirect(url_for("index"))

    if not tstr:
        flash("Please enter a time.")
        return redirect(url_for("index"))

    # Parse time formats:
    parsed = None
    for fmt in ("%H:%M", "%I:%M %p", "%I:%M%p"):
        try:
            parsed = datetime.strptime(tstr.upper(), fmt).time()
            break
        except ValueError:
            pass

    if parsed is None:
        flash("Time format wrong. Use 09:15 or 9:15 AM.")
        return redirect(url_for("index"))

    # Keep the SAME DATE, only change the time part
    if which == "in":
        old_dt = ci.time_in
        if not old_dt:
            flash("No Time In to edit.")
            return redirect(url_for("index"))
        ci.time_in = datetime.combine(old_dt.date(), parsed)

        # optional safety: if time_out exists and becomes earlier than time_in, warn
        if ci.time_out and ci.time_out < ci.time_in:
            flash("Warning: Time Out is earlier than Time In.")

    else:  # which == "out"
        if not ci.time_out:
            flash("No Time Out yet. Check out first, then edit.")
            return redirect(url_for("index"))
        old_dt = ci.time_out
        ci.time_out = datetime.combine(old_dt.date(), parsed)

        if ci.time_in and ci.time_out < ci.time_in:
            flash("Time Out cannot be earlier than Time In.")
            return redirect(url_for("index"))

    db.session.commit()
    flash("Time updated.")
    return redirect(url_for("index"))

@app.route("/version")
def version():
    last_id = db.session.query(func.max(CheckIn.id)).scalar() or 0
    last_in = db.session.query(func.max(CheckIn.time_in)).scalar()
    last_out = db.session.query(func.max(CheckIn.time_out)).scalar()

    # convert None to empty string
    return f"{last_id}|{last_in or ''}|{last_out or ''}"


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000)